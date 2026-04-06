from docx import Document
from docx.shared import Inches
from pathlib import Path
from docx.shared import Inches

"""def pngs_to_word_grid(
    image_folder='control_images/',
    out_docx="baseline_images.docx",
    rows=3,
    cols=2,
    image_width_inches=3.7
):
    image_folder = Path(image_folder)
    pngs = sorted(image_folder.glob("*.png"))
    if not pngs:
        raise FileNotFoundError(f"No PNG files found in {image_folder}")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    per_page = rows * cols
    idx = 0

    while idx < len(pngs):
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = True

        for r in range(rows):
            for c in range(cols):
                if idx >= len(pngs):
                    break

                cell = table.cell(r, c)
                p = cell.paragraphs[0]

                # Optional caption
                #p.add_run(pngs[idx].name + "\n")

                # Insert image
                p.add_run().add_picture(
                    str(pngs[idx]),
                    width=Inches(image_width_inches)
                )

                idx += 1

        if idx < len(pngs):
            doc.add_page_break()

    doc.save(out_docx)
    return out_docx

if __name__ == "__main__":
    pngs_to_word_grid()"""

# Example usage:
# pngs_to_word_grid("/home/you/pngs", "figures.docx")

from pathlib import Path
from itertools import zip_longest
from docx import Document
from docx.shared import Inches

def interleave_pngs_to_word_grid(
    baseline_folder="./baseline_images",
    control_folder="./control_images",
    out_docx="baseline_control_interleaved.docx",
    rows=3,
    cols=2,
    image_width_inches=3.7
):
    baseline_folder = Path(baseline_folder)
    control_folder = Path(control_folder)

    baseline_pngs = sorted(baseline_folder.glob("*.png"))
    control_pngs  = sorted(control_folder.glob("*.png"))

    if not baseline_pngs and not control_pngs:
        raise FileNotFoundError("No PNG files found in either folder.")

    # Interleave: baseline, control, baseline, control, ...
    ordered = []
    for b, c in zip_longest(baseline_pngs, control_pngs):
        if b is not None:
            ordered.append(b)
        if c is not None:
            ordered.append(c)

    doc = Document()

    # 0.5" margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    idx = 0
    per_page = rows * cols

    while idx < len(ordered):
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = True

        for r in range(rows):
            for c in range(cols):
                if idx >= len(ordered):
                    break

                cell = table.cell(r, c)
                p = cell.paragraphs[0]
                p.add_run().add_picture(str(ordered[idx]), width=Inches(image_width_inches))
                idx += 1

        if idx < len(ordered):
            doc.add_page_break()

    doc.save(out_docx)
    print("Saved:", Path(out_docx).resolve())


if __name__ == "__main__":
    interleave_pngs_to_word_grid()
